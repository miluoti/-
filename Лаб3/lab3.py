import os
from docx2pdf import convert
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

#Возьмем данные, полченные из программ первой и второй лабораторных:
file1 = open("output.txt", "r")
file2 = open("output2.txt", "r")
cost_t = file1.readline()
cost_t = cost_t[(cost_t.index(':') + 2):-2]
cost_s = file1.readline()
cost_s = cost_s[(cost_s.index(':') + 2):-1]
file2.readline()
cost_i = file2.readline()
cost_i = cost_i[(cost_i.index(':') + 2):-1]

#Вставим данные в таблицы файла:
doc = docx.Document('schet.docx')
tables = doc.tables

tables[0].cell(0, 0).paragraphs[0].text = 'АО "Банк Точка" Г. МОСКВА'
tables[0].cell(0, 3).paragraphs[0].text = '044184700'
tables[0].cell(1, 3).paragraphs[0].text = '74653098100000001000'
tables[0].cell(3, 0).paragraphs[0].text = 'ИНН 7722763982'
tables[0].cell(3, 1).paragraphs[0].text = 'КПП 772287204'
tables[0].cell(3, 3).paragraphs[0].text = '40700935400000001093'
tables[0].cell(4, 0).paragraphs[0].text = 'ООО "Конфетка"'

tables[1].cell(0, 0).text = 'Счет на оплату №18477 от 09 июля 2020 г.'

run = tables[1].cell(0, 0).paragraphs[0]
run.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = run.runs[0]
run.font.bold = True
run.font.size=Pt(13)

t_i = 'ООО "Конфетка", ИНН 7722763982, КПП 772287204, 109354, '
t_i += 'г. Санкт-Петербург ул. Гороховая, д. 193, корпус 5, тел: 8 (812) 192-43-56'
tables[2].cell(0, 1).text = t_i
run = tables[2].cell(0, 1).paragraphs[0].runs[0]
run.font.bold = True
t_z = 'ООО "Волна", ИНН 7736520984, КПП 777465309, 108334, '
t_z += 'г. Москва ул. Гагарина, д. 21, корпус 2, тел: 8 (495) 365-87-47'
tables[2].cell(1, 1).text = t_z
run = tables[2].cell(1, 1).paragraphs[0].runs[0]
run.font.bold = True
tables[2].cell(2, 1).text = "№20027263 от 09.07.20"
run = tables[2].cell(2, 1).paragraphs[0].runs[0]
run.font.bold = True

tables[3].cell(1, 1).text = 'Тарификация услуг типа "Телефония".'
tables[3].cell(1, 4).text = cost_t
tables[3].cell(1, 5).text = cost_t
tables[3].add_row()
tables[3].cell(2, 0).text = '2'
tables[3].cell(2, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
tables[3].cell(2, 1).text = 'Тарификация услуг типа "СМС".'
tables[3].cell(2, 4).text = cost_s
tables[3].cell(2, 5).text = cost_s
tables[3].add_row()
num = tables[3].cell(3, 0).text = '3'
tables[3].cell(3, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
tables[3].cell(3, 1).text = 'Тарификация услуг типа "Интернет".'
tables[3].cell(3, 4).text = cost_i
tables[3].cell(3, 5).text = cost_i

a = float(cost_t[:-7]) + float(cost_s[:-7]) + float(cost_i[:-7])
tables[4].cell(0,1).paragraphs[0].text = str(a)
run = tables[4].cell(0, 1).paragraphs[0].runs[0]
run.font.bold = True
tables[4].cell(1,1).paragraphs[0].text = str(round((a)/5, 2))
run = tables[4].cell(1, 1).paragraphs[0].runs[0]
run.font.bold = True
tables[4].cell(2,1).paragraphs[0].text = str(a)
run = tables[4].cell(2, 1).paragraphs[0].runs[0]
run.font.bold = True

tables[4].cell(3, 0).paragraphs[1].text = ("Всего наименований " + num + ", на сумму " + str(a) + " рубл.")

#Сумма прописью (рассчитано на четырехзначное число):
#Словари для каждого десятка:
thousand_dict = {1:"Одна тысяча", 2: "Две тысячи", 3: "Три тысячи", 4: "Четыре тысячи", 5: "Пять тысяч", 6: "Шесть тысяч", 7: "Семь тысяч", 8: "Восемь тысяч", 9: "Девять тысяч"}
hundred_dict = {1: "сто", 2: "двести", 3: "триста", 4: "четыреста", 5: "пятьсот", 6: "шестьсот", 7: "семьсот", 8: "восемьсот", 9: "девятьсот"}
decimal_dict = {2: "двадцать", 3: "тридцать", 4: "сорок", 5: "пятьдесят", 6: "шестьдесят", 7: "семьдесят", 8: "восемьдесят", 9: "девяносто"}
decimal_dict_below_20 = {10: "десять", 11: "одиннадцать", 12: "двенадцать", 13: "тринадцать", 14: "четырнадцать", 15: "пятнадцать", 16: "шестнадцать", 17: "семнадцать", 18: "восемнадцать", 19: "девятнадцать"}
digit_dict = {0: "рублей", 1: "один рубль", 2: "два рубля", 3: "три рубля", 4: "четыре рубля", 5: "пять рублей", 6: "шесть рублей", 7: "семь рублей", 8: "восемь рублей", 9: "девять рублей"}
digit_dict_cent = {0: "копеек", 1: "одна копейка", 2: "две копейки", 3: "три копейки", 4: "четыре копейки", 5: "пять копеек", 6: "шесть копеек", 7: "семь копеек", 8: "восемь копеек", 9: "девять копеек"}

integer = str(a)
fl = integer[5:]
integer = integer[:4]

answer = ''
#Распишем целую часть:
answer += thousand_dict[int(integer[0])] + ' ' + hundred_dict[int(integer[1])] + ' '
if (int(integer[2:]) > 19):
	answer += decimal_dict[int(integer[2])] + ' ' + digit_dict[int(integer[3])]
elif (int(integer[2:]) < 10):
	answer += digit_dict[int(integer[3])]
else:
	answer += decimal_dict_below_20[int(integer[2:])] + ' рублей'
answer += ' '

#Распишем дробную часть:
if (int(fl) > 19):
	answer += decimal_dict[int(fl[0])] + ' ' + digit_dict_cent[int(fl[1])]
elif (0 < int(fl) < 10):
	answer += digit_dict_cent[int(fl[1])]
elif int(fl) == 0:
	answer += ' ноль копеек'
else:
	answer += decimal_dict_below_20[int(fl)] + ' копеек'

tables[4].cell(4, 0).paragraphs[0].text = answer
run = tables[4].cell(4, 0).paragraphs[0].runs[0]
run.font.bold = True


l_str = doc.paragraphs[-1].text
l_str = l_str[:l_str.index("Б") - 9] + " Смирнов В.В.	" + l_str[l_str.index("Б"):-9] + " Кай К.Д."
doc.paragraphs[-1].text = l_str

doc.save('schet2.docx')
convert('schet2.docx', 'schet2.pdf')
